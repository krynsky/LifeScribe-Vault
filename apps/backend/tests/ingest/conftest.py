from __future__ import annotations

from pathlib import Path

import pytest


def _write_minimal_pdf(path: Path) -> None:
    content = (
        b"%PDF-1.4\n"
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
        b"2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj\n"
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]"
        b"/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
        b"4 0 obj<</Length 44>>stream\n"
        b"BT /F1 24 Tf 72 700 Td (Hello, world.) Tj ET\n"
        b"endstream endobj\n"
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n"
        b"0000000010 00000 n \n0000000053 00000 n \n"
        b"0000000100 00000 n \n0000000190 00000 n \n"
        b"0000000250 00000 n \n"
        b"trailer<</Size 6/Root 1 0 R>>\nstartxref\n310\n%%EOF\n"
    )
    path.write_bytes(content)


@pytest.fixture
def hello_pdf(tmp_path: Path) -> Path:
    p = tmp_path / "hello.pdf"
    _write_minimal_pdf(p)
    return p


@pytest.fixture
def hello_docx(tmp_path: Path) -> Path:
    from docx import Document
    doc = Document()
    doc.core_properties.title = "Hello"
    doc.core_properties.author = "Alice"
    doc.add_heading("H1 Heading", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_heading("Sub", level=2)
    doc.add_paragraph("Second paragraph.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    p = tmp_path / "h.docx"
    doc.save(str(p))
    return p
